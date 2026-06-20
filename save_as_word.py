from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word():
    doc = Document()

    # Title Style
    title = doc.add_heading('MATERYAL VE METOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    sections = [
        ("3.X. Lektin Histokimyası Görüntü Analizi ve H-Skor Hesaplamaları", 
         "Lektin uygulaması sonrası elde edilen histolojik kesitlerdeki DAB (3,3'-Diaminobenzidin) boyanma yoğunluğu, objektif ve sayısal verilere dönüştürülmek amacıyla bilgisayar destekli görüntü analiz yöntemleri ile değerlendirilmiştir."),
        
        ("3.X.1. Dijital Görüntü İşleme ve Renk Dekonvolüsyonu (Color Deconvolution)",
         "Dijital mikroskop yardımıyla alınan RGB formatındaki görüntüler, Python (v3.13) programlama dili ve 'scikit-image' kütüphanesi kullanılarak analiz edilmiştir. DAB ve Hematoksilen boyanmalarının spektral olarak birbirinden ayrılması için Ruifrok ve Johnston (2001) tarafından geliştirilen 'Renk Dekonvolüsyonu' (Color Deconvolution) algoritması uygulanmıştır. Bu yöntemle, her bir pikselin Lambert-Beer yasasına uygun olarak Optik Dansite (OD) değerleri hesaplanmış ve DAB kromojenine ait sinyal, arka plan gürültüsünden ve karşıt boyanmadan (Hematoksilen) izole edilmiştir."),
        
        ("3.X.2. H-Skor (H-Score) Analizi ve Yoğunluk Sınıflandırması",
         "İzole edilen DAB sinyallerinin kantitatif analizi için literatürle uyumlu olarak modifiye edilmiş H-Skor yöntemi kullanılmıştır. Analiz sürecinde pozitif boyanma kabulü için alt eşik değer (baseline threshold) 0.22 OD olarak belirlenmiştir. Bu değerin üzerindeki sinyaller, boyanma yoğunluklarına göre şu üç kategoriye ayrılmıştır:\n\n"
         "• Zayıf Pozitif (1+): 0.22 ≤ OD < 0.35\n"
         "• Orta Pozitif (2+): 0.35 ≤ OD < 0.50\n"
         "• Güçlü Pozitif (3+): OD ≥ 0.50\n\n"
         "Her bir görüntü alanı için ilgili yoğunluk kategorilerinin toplam alana oranı (% Pixel Area) hesaplanmış ve nihai H-Skor değeri aşağıdaki formül kullanılarak 0 ile 300 arasında bir ölçekte elde edilmiştir:\n\n"
         "H-Skor = (1 x % Zayıf Alan) + (2 x % Orta Alan) + (3 x % Güçlü Alan)\n\n"
         "Bu işlem, çalışmadaki tüm bölgeler (Alsancak, Pasaport, Urla) ve tüm lektin tipleri (WGA, SNA, MAL) için otomatize edilmiş bir Python scripti aracılığıyla standardize edilerek gerçekleştirilmiştir."),
        
        ("3.Y. İstatistiksel Analizler",
         "Elde edilen sayısal H-Skor verilerinin istatistiksel değerlendirmesinde SciPy ve Statsmodels kütüphaneleri kullanılmıştır. "
         "Verilerin normal dağılıma uygunluğu değerlendirildikten sonra, ikiden fazla grup (bölgeler arası) arasındaki farklılıkların tespiti için parametrik olmayan Kruskal-Wallis H-Testi uygulanmıştır. "
         "Kruskal-Wallis testi sonucunda anlamlı farklılık bulunan grupların (p < 0.05) ikili karşılaştırmaları, Mann-Whitney U testi kullanılarak post-hoc analiz edilmiştir.\n\n"
         "Tüm analizlerde anlamlılık düzeyi p < 0.05 olarak kabul edilmiş ve istatistiksel anlamlılık dereceleri grafikler üzerinde şu şekilde belirtilmiştir:\n"
         "* p < 0.05\n"
         "** p < 0.01\n"
         "*** p < 0.001\n\n"
         "Verilerin görselleştirilmesinde Seaborn ve Matplotlib kütüphaneleri kullanılarak, hata paylarının Standart Hata (Standard Error, SE) olarak belirtildiği bar grafikleri ve bireysel veri dağılımlarını gösteren saçılım (jitter) grafikleri oluşturulmuştur.")
    ]

    for head, text in sections:
        h = doc.add_heading(head, level=1)
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save
    file_name = "/Users/erdalbalcan/midye_lektin_olcum/materyal_metot.docx"
    doc.save(file_name)
    print(f"File saved: {file_name}")

if __name__ == "__main__":
    create_word()
